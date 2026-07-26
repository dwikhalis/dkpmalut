from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Dokumentasi Teknis dan Panduan Pengguna - BLUD Kawasan Konservasi Maluku Utara.docx"
LOGO = ROOT / "public" / "assets" / "logo_malut.png"

BLUE = "075985"
SKY = "E0F2FE"
LIGHT = "F8FAFC"
GRAY = "475569"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Klik kanan lalu pilih Update Field untuk memperbarui daftar isi."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[str]], widths=None):
    p = document.add_paragraph(caption, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        cell.text = value
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        if widths:
            cell.width = widths[idx]
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                set_cell_shading(cells[idx], LIGHT)
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if widths:
                cells[idx].width = widths[idx]
    document.add_paragraph()
    return table


def add_bullets(document: Document, items: list[str], numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        document.add_paragraph(item, style=style)


def add_note(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, SKY)
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    label = p.add_run(f"{title}: ")
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(text)
    document.add_paragraph()


def add_flow(document: Document, caption: str, steps: list[str]) -> None:
    table = document.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, step in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, BLUE if idx % 2 == 0 else "0EA5E9")
        set_cell_margins(cell, 150, 90, 150, 90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step + ("  →" if idx < len(steps) - 1 else ""))
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 255, 255)
    p = document.add_paragraph(caption, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1E293B")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 25, BLUE),
        ("Subtitle", 13, GRAY),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, "0369A1"),
        ("Heading 3", 11, GRAY),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(GRAY)

    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)


def configure_headers(document: Document) -> None:
    for section in document.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "BLUD KAWASAN KONSERVASI — DOKUMENTASI TEKNIS & PANDUAN PENGGUNA"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(GRAY)
        add_page_number(section.footer.paragraphs[0])


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.25))
    p = doc.add_paragraph("DOKUMENTASI TEKNIS APLIKASI\nDAN PANDUAN PENGGUNA", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Platform BLUD Kawasan Konservasi\nProvinsi Maluku Utara", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph("Versi dokumentasi 1.0")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(date.today().strftime("%d %B %Y"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_note(
        doc,
        "Tujuan",
        "Referensi menyeluruh bagi administrator, operator tiket, pengelola konten, pengguna terdaftar, dan tim teknis yang mengoperasikan aplikasi.",
    )
    p = doc.add_paragraph("Disusun berdasarkan kondisi source code pada tanggal penerbitan.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    doc.add_page_break()

    doc.add_heading("Informasi Dokumen", 1)
    add_table(
        doc,
        "Tabel 1. Informasi dokumen",
        ["Item", "Keterangan"],
        [
            ["Judul", "Dokumentasi Teknis dan Panduan Pengguna — BLUD Kawasan Konservasi Maluku Utara"],
            ["Jenis sistem", "Aplikasi web responsif untuk informasi publik, data, peta kawasan konservasi, dan layanan tiket"],
            ["Audiens", "Publik, pengguna terdaftar, administrator, petugas pemindai tiket, dan pengelola teknis"],
            ["Bahasa aplikasi", "Bahasa Indonesia dan Bahasa Inggris (konten bergantung kelengkapan konfigurasi CMS)"],
            ["Sumber acuan", "Source code aplikasi, konfigurasi deployment, migrasi SQL, dan runbook produksi"],
            ["Catatan", "Struktur database produksi wajib diverifikasi terhadap proyek Supabase aktif sebelum migrasi atau pemulihan."],
        ],
    )
    doc.add_heading("Daftar Isi", 1)
    add_toc(doc)
    doc.add_page_break()

    # I
    doc.add_heading("I. Pendahuluan", 1)
    doc.add_heading("1.1 Tujuan Dokumen", 2)
    doc.add_paragraph(
        "Dokumen ini menjelaskan fungsi aplikasi, arsitektur teknis, komponen utama, penyimpanan data, keamanan, integrasi pembayaran, deployment, pemeliharaan, serta langkah penggunaan untuk publik dan administrator. Dokumen dirancang sebagai laporan teknis sekaligus manual operasional."
    )
    doc.add_heading("1.2 Latar Belakang", 2)
    doc.add_paragraph(
        "Platform BLUD Kawasan Konservasi Provinsi Maluku Utara menyediakan satu pintu untuk informasi organisasi, kawasan konservasi, dataset kelautan dan perikanan, peta interaktif, berita, galeri, regulasi, tarif, kontak, serta pembelian dan validasi tiket. Konten publik dikelola melalui dashboard sehingga pembaruan rutin tidak mengharuskan perubahan source code."
    )
    doc.add_heading("1.3 Sasaran Pengguna", 2)
    add_bullets(doc, [
        "Pengunjung publik yang melihat informasi, data, peta, berita, galeri, tarif, dan regulasi.",
        "Pengguna terdaftar yang mengelola profil, melihat riwayat tiket, atau memperoleh akses data tertentu.",
        "Administrator yang mengelola konten, dataset, kawasan, pengguna, pesan, tarif, dan aktivitas.",
        "Petugas lapangan yang memvalidasi tiket menggunakan pemindai QR.",
        "Pengelola teknis yang memelihara aplikasi, database, integrasi, keamanan, backup, dan deployment.",
    ])

    # II
    doc.add_heading("II. Ringkasan Platform", 1)
    doc.add_heading("2.1 Identitas dan Ruang Lingkup", 2)
    add_table(doc, "Tabel 2. Identitas platform", ["Komponen", "Keterangan"], [
        ["Nama tampilan", "BLUD Kawasan Konservasi — Provinsi Maluku Utara"],
        ["Jenis aplikasi", "Web application responsif dengan rendering Next.js"],
        ["Frontend", "React 19, Next.js 16 App Router, TypeScript, Tailwind CSS"],
        ["Backend dan data", "Next.js Route Handlers, Supabase/PostgreSQL, Supabase Auth dan Storage"],
        ["Integrasi", "Midtrans Snap, SMTP, Cloudflare Turnstile, OpenStreetMap/Leaflet, Vercel Analytics"],
        ["Perangkat", "Browser desktop dan mobile; kamera diperlukan untuk pemindaian QR"],
        ["Deployment acuan", "Vercel dengan cron harian dan environment variables"],
    ])
    doc.add_heading("2.2 Modul Utama", 2)
    add_table(doc, "Tabel 3. Modul aplikasi", ["Modul", "Kemampuan utama"], [
        ["Situs publik", "Beranda, organisasi, staf, kawasan konservasi, berita, galeri, regulasi, FAQ, kontak, kebijakan, dan aksesibilitas."],
        ["Data dan visualisasi", "Daftar dataset, tabel dinamis, filter, pengurutan, grafik, CSV, dan konfigurasi publikasi."],
        ["Peta", "Dataset peta, layer GeoJSON, legenda, pola isi, marker, overlay, dan dokumen pendukung."],
        ["Akun", "Pendaftaran, konfirmasi email, login, reset password, profil, dan proteksi akses."],
        ["CMS", "Teks, ikon, gambar, visibilitas komponen, pratinjau, dan locale ID/EN."],
        ["Tiket", "Persiapan pesanan, perhitungan biaya, pembayaran Midtrans, PDF/QR, email, riwayat, verifikasi, dan scan."],
        ["Administrasi", "Statistik, dataset, pengguna, pesan, kawasan, tarif, konten, konfigurasi tabel, dan audit log."],
    ])

    # III
    doc.add_heading("III. Teknologi dan Dependensi", 1)
    add_table(doc, "Tabel 4. Teknologi yang digunakan", ["Lapisan", "Teknologi", "Fungsi"], [
        ["Framework", "Next.js 16.2 / React 19", "Routing, rendering server/client, route handler, optimasi gambar, dan build."],
        ["Bahasa", "TypeScript 6", "Type safety untuk komponen, API, store, dan model data."],
        ["UI", "Tailwind CSS 4", "Layout responsif, tema visual, dan state interaksi."],
        ["State", "Zustand", "Status autentikasi, locale, konten admin, dan checkout tiket."],
        ["Database", "Supabase PostgreSQL", "Data relasional, JSONB konfigurasi, RLS, fungsi database, trigger, dan realtime."],
        ["Penyimpanan", "Supabase Storage", "Gambar, dokumen, dan berkas GeoJSON."],
        ["Peta", "Leaflet / React Leaflet", "Peta interaktif, layer, marker, popup, dan GeoJSON."],
        ["Grafik", "Chart.js", "Visualisasi dataset dan label data."],
        ["Pembayaran", "Midtrans Snap", "Pembuatan transaksi, pembayaran, status, pembatalan, dan webhook."],
        ["Dokumen", "pdf-lib / QRCode", "Pembuatan PDF tiket dan kode QR."],
        ["Email", "Nodemailer / SMTP", "Notifikasi kontak dan pengiriman tiket."],
        ["Keamanan", "Turnstile + CSP + rate limit", "Perlindungan bot, pembatasan request, dan kebijakan browser."],
        ["Hosting", "Vercel", "Deployment Next.js, environment, analytics, cron, dan rollback."],
    ])
    doc.add_heading("3.1 Perintah Pengembangan", 2)
    add_table(doc, "Tabel 5. Perintah proyek", ["Perintah", "Kegunaan"], [
        ["npm run dev", "Menjalankan development server dengan Turbopack."],
        ["npm run lint", "Menjalankan ESLint dan menolak warning."],
        ["npm test", "Menjalankan pengujian keamanan pembayaran."],
        ["npm run build", "Build produksi serta pemeriksaan TypeScript."],
        ["npm run check", "Menjalankan lint, test, dan build secara berurutan."],
        ["npm run validate:production-env", "Memvalidasi environment variable wajib untuk produksi."],
    ])

    # IV
    doc.add_heading("IV. Arsitektur Sistem", 1)
    doc.add_heading("4.1 Gambaran Umum", 2)
    doc.add_paragraph(
        "Aplikasi memakai arsitektur web berlapis. Browser merender UI React. Halaman server dan client berkomunikasi dengan Supabase menggunakan anon key serta session pengguna. Operasi sensitif—pembayaran, webhook, rate limit, audit, dan pekerjaan email—dilakukan melalui route handler server menggunakan kredensial privat. Midtrans mengirim notifikasi pembayaran ke endpoint server, sedangkan SMTP mengirim tiket dan notifikasi."
    )
    add_flow(doc, "Gambar 1. Alur arsitektur tingkat tinggi", [
        "Browser",
        "Next.js / Vercel",
        "Supabase",
        "Midtrans",
        "SMTP",
    ])
    doc.add_heading("4.2 Lapisan Aplikasi", 2)
    add_table(doc, "Tabel 6. Lapisan aplikasi", ["Lapisan", "Lokasi utama", "Tanggung jawab"], [
        ["Routing/UI", "app/**/page.tsx, app/components", "Halaman publik, dashboard, formulir, tabel, peta, dan tiket."],
        ["State dan auth", "app/Stores, app/Auth", "Session, role, locale, checkout, redirect, dan proteksi halaman."],
        ["API server", "app/api/**/route.ts", "Validasi request, pembayaran, webhook, tiket, scan, kontak, audit, dan cron."],
        ["Domain/utilitas", "lib/**", "Supabase helper, email, PDF, keamanan pembayaran, peta, konfigurasi, dan normalisasi."],
        ["Database", "Supabase PostgreSQL", "Data, constraints, RLS, trigger audit, realtime, dan fungsi rate limit."],
        ["Object storage", "Supabase Storage", "Bucket images, geojsons, dan documents."],
    ])
    doc.add_heading("4.3 Alur Data Publik", 2)
    add_bullets(doc, [
        "Pengunjung membuka route publik; Next.js memuat shell, metadata, serta data yang diperlukan.",
        "Data terstruktur dibaca dari Supabase dan ditampilkan sebagai kartu, tabel, grafik, atau peta.",
        "Locale store memilih konten Bahasa Indonesia atau Inggris; fallback digunakan saat terjemahan belum tersedia.",
        "Filter dan query URL menjaga tampilan data dapat dibagikan dan dinavigasi.",
        "Berkas publik disajikan dari folder public atau URL Supabase Storage.",
    ], numbered=True)
    doc.add_heading("4.4 Alur Administrasi", 2)
    add_bullets(doc, [
        "Supabase Auth memverifikasi session.",
        "AuthProvider memuat profil dan role dari tabel users.",
        "AuthProtect membatasi halaman terdaftar; AuthAdminAccess mensyaratkan role admin.",
        "Perubahan dilakukan melalui helper Supabase atau endpoint server sesuai tingkat sensitivitas.",
        "RLS dan trigger database menjadi lapisan otorisasi serta audit tambahan.",
    ], numbered=True)

    # V
    doc.add_heading("V. Struktur Aplikasi dan Routing", 1)
    add_table(doc, "Tabel 7. Route publik utama", ["Route", "Fungsi"], [
        ["/", "Beranda dan ringkasan layanan."],
        ["/organisasi, /organisasi/staff", "Profil organisasi dan daftar staf."],
        ["/explore, /explore/[slug]", "Daftar dan detail kawasan konservasi."],
        ["/explore/[slug]/peta", "Peta detail kawasan konservasi."],
        ["/data, /data/[slug]", "Katalog dan publikasi dataset/peta."],
        ["/berita, /berita/[id]", "Daftar dan detail berita/artikel."],
        ["/galeri, /peraturan, /faq", "Galeri, regulasi, dan pertanyaan umum."],
        ["/informasi-tarif", "Tarif kawasan dan komponen biaya aktif."],
        ["/kontak", "Formulir kontak dan status pengiriman."],
        ["/payment", "Checkout dan pembayaran tiket."],
        ["/ticket/verify/[ticketCode]", "Verifikasi informasi tiket melalui kode."],
        ["/masuk, /daftar, /reset-password", "Autentikasi dan pemulihan akun."],
    ])
    add_table(doc, "Tabel 8. Endpoint server", ["Endpoint", "Metode/fungsi"], [
        ["/api/contact", "Menyimpan pesan, memeriksa pengguna platform, dan mengirim email notifikasi."],
        ["/api/tickets/prepare", "Validasi checkout, harga authoritative, dan pembuatan booking pending."],
        ["/api/midtrans/create-transaction", "Membuat Snap transaction berdasarkan booking valid."],
        ["/api/midtrans/notification", "Memverifikasi signature dan memproses perubahan status pembayaran."],
        ["/api/tickets/status", "Menyelaraskan status pembayaran dan status email tiket."],
        ["/api/midtrans/reset-transaction", "Membatalkan/mengganti sesi pembayaran pending."],
        ["/api/tickets/history", "Mengambil riwayat tiket milik pengguna terautentikasi."],
        ["/api/tickets/scan", "Memvalidasi dan menandai pemakaian tiket oleh admin."],
        ["/api/activity-logs", "Menyediakan audit log hanya untuk admin."],
        ["/api/cron/cleanup-drafts", "Menghapus draft dataset/peta kedaluwarsa dengan CRON_SECRET."],
    ])

    # VI
    doc.add_heading("VI. Modul dan Fitur Utama", 1)
    modules = [
        ("6.1 Konten Publik dan Navigasi", "Navbar responsif menyediakan menu Tentang, berita, galeri, data, kontak, pemilih bahasa, dan akun. Halaman CMS menggunakan nilai fallback agar aplikasi tetap dapat dirender apabila konfigurasi tertentu belum tersedia."),
        ("6.2 Kawasan Konservasi", "Data kawasan mencakup slug, nama ringkas/resmi, kategori, lokasi, luas, ringkasan, ekosistem, fitur utama, zonasi, dokumen, gambar peta, urutan tampilan, dan status aktif. Konten terlokalisasi disimpan dalam JSONB."),
        ("6.3 Dataset Dinamis", "Metadata dataset menentukan label, pemilik, status publikasi, konfigurasi kolom/filter, grafik, dan tabel fisik sumber. Admin dapat membuat, mengimpor CSV, mengedit, memfilter, mempublikasikan, atau menghapus dataset."),
        ("6.4 Peta dan GeoJSON", "Peta dibangun dari map_datasets, map_layers, map_legend_items, serta file pada bucket geojsons. Editor mendukung konfigurasi layer, legenda, warna, pola isi, marker, overlay, popup, visibilitas, dan pengurutan."),
        ("6.5 Berita, Galeri, dan Staf", "Konten dikelola melalui komponen generik tambah/edit/list. Pilihan tag, divisi, jabatan, dan gender berasal dari table_config agar label ID/EN dapat diubah tanpa migrasi setiap record."),
        ("6.6 App CMS", "Tabel app_cms menyimpan target, component, type, locale, value, dan status visibilitas. Admin dapat mencari, memfilter, mengubah teks/gambar/ikon, serta membuka mode pratinjau."),
        ("6.7 Pesan Kontak", "Pesan disimpan dengan status new/old. Pengiriman email memakai status not_attempted, pending, sent, atau failed, disertai waktu kirim dan error terkontrol. Dashboard menerima perubahan melalui Supabase Realtime."),
        ("6.8 Tiket dan Tarif", "Admin mengelola kawasan yang dapat dipesan dan item biaya. Checkout menghitung biaya pada server, menyimpan visitor serta kawasan terpilih, membuat pembayaran, dan meneruskan transaksi ke Midtrans."),
        ("6.9 PDF, Email, dan QR", "Setelah pembayaran valid, aplikasi menetapkan masa berlaku, membuat PDF tiket berisi QR, dan mengirimnya melalui SMTP. Claim status email mencegah proses bersamaan dan mendukung retry saat gagal."),
        ("6.10 Audit Aktivitas", "Trigger database merekam INSERT, UPDATE, dan DELETE untuk entitas administratif. Metadata menyimpan label dan field berubah, tetapi mengecualikan token, password, kode tiket, dan signature."),
    ]
    for heading, body in modules:
        doc.add_heading(heading, 2)
        doc.add_paragraph(body)

    # VII
    doc.add_heading("VII. Database dan Penyimpanan", 1)
    doc.add_heading("7.1 Entitas Utama", 2)
    add_table(doc, "Tabel 9. Tabel database utama", ["Tabel", "Fungsi", "Catatan penting"], [
        ["users", "Profil pengguna dan role.", "Terhubung ke Supabase Auth; role dibaca oleh AuthProvider."],
        ["app_cms", "Konten dan visibilitas UI per locale.", "Nilai dapat berupa teks, ikon, atau path gambar."],
        ["staff / news / gallery", "Konten organisasi dan publikasi.", "Pilihan terkontrol oleh table_config."],
        ["table_config", "Konfigurasi dinamis staf/berita/galeri.", "Publik dapat membaca; admin mengelola melalui RPC terlindungi."],
        ["datasets", "Metadata dataset dinamis.", "Berelasi dengan tabel sumber dan konfigurasi JSON."],
        ["map_datasets", "Metadata publikasi peta.", "Mengikat layer dan konfigurasi tampilan."],
        ["map_layers", "Definisi layer peta.", "Merujuk file GeoJSON dan urutan/visibilitas."],
        ["map_legend_items", "Legenda dan gaya layer.", "Mendukung pola none, diagonal, crosshatch, horizontal, vertical, dots."],
        ["conservation_areas", "Konten kawasan dan konfigurasi tiket.", "Konten bilingual JSONB serta luas dan urutan."],
        ["ticket_charge_items", "Komponen tarif tiket.", "Server memakai nilai aktif ketika menghitung total."],
        ["payments", "Booking dan transaksi pembayaran.", "Status, order id, token status, waktu bayar, QR, pemakaian, email."],
        ["ticket_visitors", "Identitas pengunjung per booking.", "Jangan dicatat ke log eksternal."],
        ["payment_conservation_areas", "Kawasan yang terhubung ke booking.", "Snapshot pilihan dan nilai terkait."],
        ["messages", "Pesan formulir kontak.", "Status baca dan status pengiriman email."],
        ["activity_logs", "Jejak perubahan admin.", "Read melalui endpoint admin; tabel tidak terbuka langsung."],
    ])
    doc.add_heading("7.2 Status Penting", 2)
    add_table(doc, "Tabel 10. Status domain", ["Domain", "Nilai", "Makna"], [
        ["Pesan", "new / old", "Pesan baru atau telah ditandai lama/dibaca."],
        ["Email kontak", "not_attempted / pending / sent / failed", "Siklus pengiriman notifikasi kontak."],
        ["Pembayaran", "pending / paid / failed / expired / cancelled / challenge", "Hasil normalisasi status Midtrans."],
        ["Email tiket", "not_sent / sending / sent / failed", "Siklus claim dan pengiriman lampiran tiket."],
        ["Publikasi data", "requested / approved / rejected", "Tahap review dataset sebelum ditampilkan."],
    ])
    doc.add_heading("7.3 Storage", 2)
    add_table(doc, "Tabel 11. Bucket Supabase Storage", ["Bucket", "Isi", "Penggunaan"], [
        ["images", "Logo, foto, hero, thumbnail, dan gambar peta.", "CMS, berita, galeri, staf, profil, kawasan, dataset."],
        ["geojsons", "File data spasial.", "Diunduh dan diparsing oleh modul peta."],
        ["documents", "Dokumen kawasan/regulasi.", "URL publik untuk unduhan atau referensi."],
    ])
    add_note(doc, "Penting", "Migrasi SQL dalam repository bersifat inkremental dan bukan satu-satunya definisi baseline database. Sebelum pemulihan atau pembuatan environment baru, ekspor schema produksi dengan Supabase CLI dan simpan sebagai baseline terversi.")

    # VIII
    doc.add_heading("VIII. Autentikasi, Otorisasi, dan Keamanan", 1)
    doc.add_heading("8.1 Autentikasi dan Role", 2)
    add_table(doc, "Tabel 12. Hak akses", ["Peran", "Akses umum"], [
        ["Publik", "Membaca halaman publik, melihat data/peta, tarif, dan memulai pembelian tiket."],
        ["User", "Fitur publik ditambah profil dan riwayat tiket milik sendiri."],
        ["Admin", "Dashboard, manajemen konten/data/pengguna/tarif, scan tiket, audit, dan konfigurasi."],
    ])
    doc.add_paragraph(
        "Supabase Auth menyimpan kredensial dan session. Tabel users menyimpan profil aplikasi serta role. Proteksi UI membantu pengalaman pengguna, tetapi keputusan keamanan utama harus tetap ditegakkan oleh RLS, fungsi database, dan pemeriksaan server."
    )
    doc.add_heading("8.2 Kontrol Keamanan", 2)
    add_bullets(doc, [
        "Cloudflare Turnstile pada login dan pendaftaran untuk mengurangi otomasi bot.",
        "Rate limit server menggunakan fungsi consume_api_rate_limit dan salt privat.",
        "Signature Midtrans diverifikasi sebelum status pembayaran diperbarui.",
        "Total pembayaran dihitung ulang dari data tarif server; nilai dari browser tidak dipercaya.",
        "Status token tiket digunakan untuk pemeriksaan tanpa mengekspos kredensial pembayaran.",
        "Pemindaian tiket memerlukan bearer token pengguna dan role admin.",
        "CSP membatasi script, frame, image, connect, dan sumber eksternal.",
        "Header HSTS produksi, nosniff, SAMEORIGIN, Referrer-Policy, Permissions-Policy, dan COOP.",
        "Audit trigger menyaring field sensitif dari metadata.",
        "CRON_SECRET dibandingkan secara aman untuk endpoint cleanup.",
    ])
    doc.add_heading("8.3 Data Sensitif", 2)
    add_note(doc, "Larangan", "Jangan menyimpan atau meneruskan request body pembayaran, identitas pengunjung, SMTP password, service-role key, status token, kode tiket, maupun payload notifikasi mentah ke log atau layanan monitoring pihak ketiga.")

    # IX
    doc.add_heading("IX. Integrasi Tiket dan Pembayaran", 1)
    doc.add_heading("9.1 Alur Checkout", 2)
    add_flow(doc, "Gambar 2. Alur transaksi tiket", [
        "Isi data",
        "Prepare",
        "Booking pending",
        "Midtrans Snap",
        "Webhook/status",
        "PDF & email",
    ])
    add_bullets(doc, [
        "Pengunjung memilih kawasan, tanggal/keperluan, dan mengisi data visitor.",
        "Endpoint prepare memvalidasi format, batas jumlah, kawasan aktif, dan tarif dari database.",
        "Server menyimpan payment pending, ticket_visitors, dan payment_conservation_areas.",
        "Endpoint create-transaction membuat Snap token untuk booking yang masih pending.",
        "Midtrans memproses metode pembayaran dan mengirim notification webhook.",
        "Server memverifikasi signature serta kesesuaian gross amount sebelum memperbarui status.",
        "Status paid menetapkan masa berlaku tiket, lalu job menghasilkan PDF/QR dan mengirim email.",
        "Halaman finish/status melakukan rekonsiliasi bila webhook terlambat dan menampilkan hasil aman.",
    ], numbered=True)
    doc.add_heading("9.2 Normalisasi Status Midtrans", 2)
    add_table(doc, "Tabel 13. Pemetaan status pembayaran", ["Status Midtrans", "Status aplikasi"], [
        ["settlement", "paid"],
        ["capture (fraud_status=accept)", "paid"],
        ["pending", "pending"],
        ["expire", "expired"],
        ["cancel", "cancelled"],
        ["deny / failure", "failed"],
        ["challenge", "challenge"],
    ])
    doc.add_heading("9.3 Validasi dan Pemakaian Tiket", 2)
    doc.add_paragraph(
        "QR mengarah ke identitas tiket yang dapat diverifikasi. Petugas admin menggunakan halaman scanner, mengizinkan kamera, memindai kode, dan memeriksa status paid, tanggal kedaluwarsa, serta used_at. Pemakaian pertama menandai tiket; scan berikutnya dilaporkan sebagai sudah digunakan agar tiket tidak dipakai berulang."
    )

    # X
    doc.add_heading("X. Panduan Pengguna Publik", 1)
    procedures = [
        ("10.1 Mengubah Bahasa", ["Pada Navbar pilih ikon bahasa.", "Pilih Bahasa atau English.", "Tunggu konten dimuat ulang. Konten tanpa terjemahan dapat memakai fallback."]),
        ("10.2 Menjelajahi Kawasan", ["Buka menu Tentang lalu Kawasan.", "Pilih salah satu kartu kawasan.", "Baca ringkasan, ekosistem, fitur utama, zonasi, dan dokumen.", "Buka peta kawasan bila tersedia."]),
        ("10.3 Menggunakan Data", ["Buka menu Data.", "Pilih dataset atau peta.", "Gunakan pencarian, filter, pengurutan, atau grafik.", "Gunakan download hanya bila fitur dan hak akses mengizinkan."]),
        ("10.4 Membuat Akun", ["Buka Masuk lalu pilih pendaftaran.", "Isi data yang diminta dan selesaikan Turnstile.", "Kirim formulir.", "Buka email konfirmasi dan ikuti tautan.", "Masuk setelah akun terverifikasi."]),
        ("10.5 Membeli Tiket", ["Buka Pembayaran/Beli Tiket.", "Pilih kawasan dan isi seluruh data pengunjung dengan benar.", "Tinjau ringkasan harga.", "Lanjutkan ke Midtrans dan pilih metode pembayaran.", "Jangan menutup halaman sebelum hasil pembayaran ditampilkan.", "Simpan email dan PDF tiket; periksa folder spam jika belum terlihat."]),
        ("10.6 Melihat Riwayat Tiket", ["Masuk ke akun.", "Buka menu Akun lalu Tiket Saya.", "Lihat status, kawasan, masa berlaku, dan detail transaksi milik akun."]),
        ("10.7 Mengirim Pesan", ["Buka Kontak.", "Isi nama, email, telepon bila diperlukan, dan pesan.", "Kirim formulir dan tunggu notifikasi.", "Jika email notifikasi gagal, pesan tetap dapat tersimpan dan terlihat oleh admin."]),
    ]
    for heading, steps in procedures:
        doc.add_heading(heading, 2)
        add_bullets(doc, steps, numbered=True)

    # XI
    doc.add_heading("XI. Panduan Administrator", 1)
    admin_procedures = [
        ("11.1 Membuka Dashboard", ["Masuk menggunakan akun ber-role admin.", "Pada Navbar pilih Akun lalu Dashboard.", "Gunakan sidebar untuk membuka modul. Jika akses ditolak, periksa session dan role pada tabel users."]),
        ("11.2 Mengelola CMS", ["Buka App CMS.", "Filter berdasarkan component, locale, type, target, atau pencarian.", "Pilih item dan ubah value atau visibilitas.", "Untuk gambar, unggah aset lalu simpan path yang dihasilkan.", "Gunakan preview dan periksa desktop serta mobile sebelum publikasi."]),
        ("11.3 Mengelola Kawasan", ["Buka menu Kawasan Konservasi.", "Tambah atau edit identitas bilingual, luas, urutan, dan status aktif.", "Lengkapi ekosistem, fitur utama, zonasi, dokumen, dan gambar peta.", "Simpan lalu periksa halaman Explore dan peta publik."]),
        ("11.4 Mengelola Dataset", ["Buka Data dan pilih tambah/edit dataset.", "Tentukan jenis tabel atau peta dan metadata publikasi.", "Untuk CSV, unggah file, tinjau header dan preview, lalu pilih konfigurasi kolom/filter.", "Simpan sebagai draft selama penyusunan.", "Uji tabel, grafik, filter, dan peta sebelum mengubah status publikasi."]),
        ("11.5 Mengelola Staf/Berita/Galeri", ["Pilih modul konten.", "Gunakan List untuk mengedit/menghapus atau Tambah untuk record baru.", "Isi semua field dan unggah gambar yang sesuai.", "Simpan lalu periksa halaman publik.", "Gunakan Table Config untuk mengelola pilihan tag/divisi/jabatan/gender."]),
        ("11.6 Mengelola Tarif", ["Buka konfigurasi Ticketing.", "Aktifkan/nonaktifkan kawasan yang menerima tiket.", "Tambah atau edit item biaya dan nominal.", "Pastikan tarif efektif sebelum transaksi diuji.", "Lakukan transaksi sandbox setelah perubahan signifikan."]),
        ("11.7 Menangani Pesan", ["Buka menu Pesan.", "Kelompokkan pesan baru dan lama.", "Buka pesan untuk detail kontak serta isi.", "Periksa badge status email; failed menampilkan detail yang aman.", "Ubah status atau hapus hanya setelah kebutuhan retensi dipertimbangkan."]),
        ("11.8 Mengelola Pengguna", ["Buka Pengguna.", "Cari akun dan periksa identitas/profil.", "Ubah role hanya sesuai otorisasi organisasi.", "Verifikasi bahwa admin baru memahami akses data dan tanggung jawab keamanan."]),
        ("11.9 Memindai Tiket", ["Buka halaman scanner admin.", "Izinkan kamera atau masukkan kode bila UI menyediakan alternatif.", "Arahkan QR ke kamera dan tunggu hasil.", "Cocokkan informasi pengunjung/kawasan.", "Konfirmasi validasi; jangan menerima tiket expired, cancelled, unpaid, atau sudah digunakan."]),
        ("11.10 Meninjau Audit Log", ["Buka Aktivitas.", "Filter actor, entity, action, atau waktu.", "Tinjau changed_fields untuk perubahan penting.", "Eskalasi aktivitas mencurigakan dan rotasi kredensial bila ada indikasi kompromi."]),
    ]
    for heading, steps in admin_procedures:
        doc.add_heading(heading, 2)
        add_bullets(doc, steps, numbered=True)
    add_note(doc, "Peringatan", "Penghapusan record atau file dapat tidak dapat dipulihkan tanpa backup. Pastikan target benar, lakukan backup, dan gunakan proses persetujuan untuk data produksi.")

    # XII
    doc.add_heading("XII. Konfigurasi Environment", 1)
    add_table(doc, "Tabel 14. Environment variables", ["Variabel", "Sifat", "Fungsi"], [
        ["NEXT_PUBLIC_SUPABASE_URL", "Publik", "URL proyek Supabase."],
        ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Publik", "Anon key yang dibatasi RLS."],
        ["SUPABASE_SERVICE_ROLE_KEY", "Rahasia", "Operasi server administratif; tidak boleh masuk client bundle."],
        ["NEXT_PUBLIC_SITE_URL / NEXT_PUBLIC_BASE_URL", "Publik", "Base URL callback, tiket, dan tautan absolut."],
        ["NEXT_PUBLIC_MIDTRANS_CLIENT_KEY", "Publik", "Client key Snap."],
        ["MIDTRANS_SERVER_KEY", "Rahasia", "Otentikasi API dan verifikasi Midtrans."],
        ["MIDTRANS_IS_PRODUCTION", "Server", "Memilih sandbox atau production."],
        ["NEXT_PUBLIC_TURNSTILE_SITE_KEY", "Publik", "Site key widget Turnstile."],
        ["TURNSTILE_SECRET_KEY", "Rahasia", "Validasi token Turnstile pada server/provider auth."],
        ["TURNSTILE_EXPECTED_HOSTNAME", "Server", "Hostname yang diperbolehkan untuk token."],
        ["SMTP_HOST / SMTP_PORT / SMTP_USER / SMTP_PASSWORD", "Rahasia", "Transport pengiriman email."],
        ["SMTP_TLS_CA_FILE / SMTP_TLS_REJECT_UNAUTHORIZED", "Server", "Pengaturan trust TLS SMTP; jangan menonaktifkan verifikasi di produksi."],
        ["SMTP_ADMIN_EMAIL / SMTP_ADMIN_NAME", "Server", "Tujuan notifikasi admin."],
        ["CONTACT_FROM_EMAIL", "Server", "Alamat From notifikasi kontak."],
        ["SMTP_TICKET_FROM_EMAIL / TICKET_FROM_EMAIL", "Server", "Alamat From pengiriman tiket."],
        ["ORG_EMAIL", "Server", "Alamat organisasi untuk fallback/komunikasi."],
        ["RATE_LIMIT_SALT", "Rahasia", "Salt panjang acak untuk identitas rate limit."],
        ["CRON_SECRET", "Rahasia", "Bearer token untuk cleanup draft terjadwal."],
    ])
    add_note(doc, "Aturan", "Jangan commit file .env, service-role key, server key Midtrans, SMTP password, CRON_SECRET, atau RATE_LIMIT_SALT. Gunakan secret manager platform deployment dan rotasi setelah dugaan kebocoran.")

    # XIII
    doc.add_heading("XIII. Deployment dan Operasi Produksi", 1)
    doc.add_heading("13.1 Persiapan Deployment", 2)
    add_bullets(doc, [
        "Siapkan proyek Supabase untuk environment target dan terapkan seluruh migrasi yang relevan.",
        "Pastikan baseline schema mencakup fungsi consume_api_rate_limit yang dirujuk aplikasi.",
        "Isi environment variables pada Vercel; gunakan kredensial sandbox untuk staging.",
        "Jalankan npm run validate:production-env dan npm run check.",
        "Daftarkan URL HTTPS /api/midtrans/notification pada dashboard Midtrans.",
        "Pastikan domain callback Supabase Auth dan hostname Turnstile sesuai.",
        "Pastikan bucket Storage, policy, RLS, index, trigger audit, dan Supabase Realtime terkonfigurasi.",
    ], numbered=True)
    doc.add_heading("13.2 Cron dan Retensi Draft", 2)
    doc.add_paragraph(
        "vercel.json menjadwalkan /api/cron/cleanup-drafts pada 02:15 UTC setiap hari. Vercel mengirim Authorization Bearer menggunakan CRON_SECRET. Pada platform lain, buat jadwal GET atau POST terautentikasi sekali sehari. Pantau kegagalan job agar draft lama tidak menumpuk."
    )
    doc.add_heading("13.3 Prosedur Rilis", 2)
    add_bullets(doc, [
        "Deploy ke staging dengan Midtrans sandbox dan Supabase staging.",
        "Uji skenario berhasil, pending, gagal, webhook duplikat, tiket expired, dan double scan.",
        "Periksa CSP untuk Midtrans Snap, Turnstile, peta, Supabase realtime, dan analytics.",
        "Deploy production dan lakukan satu transaksi bernilai rendah.",
        "Verifikasi email, PDF, QR, scan, riwayat, payment row, dan webhook.",
        "Pertahankan deployment sebelumnya untuk rollback cepat.",
    ], numbered=True)
    doc.add_heading("13.4 Backup dan Pemulihan", 2)
    add_bullets(doc, [
        "Aktifkan point-in-time recovery atau backup harian Supabase sesuai paket.",
        "Backup database dan inventaris object storage; database saja tidak memulihkan file.",
        "Uji restore ke proyek terpisah sebelum peluncuran dan minimal setiap kuartal.",
        "Catat project ID, pemilik secret, DNS, Midtrans merchant, dan akun SMTP di repositori kredensial organisasi.",
        "Setelah insiden, rotasi service-role, Midtrans, SMTP, cron, rate-limit salt, dan token terkait.",
    ])

    # XIV
    doc.add_heading("XIV. Monitoring dan Penanganan Masalah", 1)
    add_table(doc, "Tabel 15. Panduan troubleshooting", ["Gejala", "Pemeriksaan", "Tindakan"], [
        ["Login/daftar gagal", "Turnstile, callback URL, rate limit, Supabase Auth, email confirmation.", "Perbaiki konfigurasi provider/hostname; jangan tampilkan error internal mentah."],
        ["Data tidak muncul", "Status publikasi, RLS, nama tabel, konfigurasi dataset, network browser.", "Validasi metadata dan policy; uji dengan role yang sama."],
        ["Peta kosong", "File GeoJSON, bucket policy, map layer, geometry, CSP tile source.", "Unduh file langsung, validasi GeoJSON, lalu periksa layer/legend."],
        ["Pembayaran tetap pending", "Webhook Midtrans, signature, order id, status endpoint.", "Cek log teredaksi dan lakukan rekonsiliasi status; jangan edit paid secara manual tanpa bukti."],
        ["Email tiket gagal", "ticket_email_status, SMTP/TLS, from address, attachment generation.", "Perbaiki SMTP lalu gunakan alur retry yang tersedia."],
        ["QR ditolak", "Status paid, expires_at, used_at, role scanner, token.", "Pastikan tiket asli dan belum dipakai; eskalasi jika data tidak konsisten."],
        ["Kontak tidak terkirim", "Row messages dan email_delivery_status.", "Pesan dapat tetap tersimpan; tindak lanjuti dari dashboard dan perbaiki SMTP."],
        ["Build gagal font", "Akses Google Fonts ketika next/font melakukan build.", "Izinkan akses build atau migrasikan font menjadi aset lokal."],
        ["Cron gagal", "Authorization Bearer dan CRON_SECRET.", "Samakan secret, uji endpoint aman, lalu pantau jadwal berikutnya."],
    ])
    doc.add_heading("14.1 Alert yang Disarankan", 2)
    add_bullets(doc, [
        "Rasio 5xx lebih dari 1% selama lima menit.",
        "503 berulang dari backend rate limit.",
        "Kegagalan endpoint cleanup draft.",
        "Webhook Midtrans gagal atau nominal pembayaran tidak sesuai.",
        "ticket_email_status menetap failed atau sending.",
        "Uptime gagal untuk beranda dan alur tiket publik.",
    ])

    # XV
    doc.add_heading("XV. Pemeliharaan dan Kontrol Perubahan", 1)
    add_bullets(doc, [
        "Gunakan branch dan review untuk perubahan aplikasi serta migrasi SQL.",
        "Jalankan lint, test, dan build sebelum merge/deploy.",
        "Buat migrasi yang idempotent bila memungkinkan dan uji pada salinan staging.",
        "Jangan mengubah status pembayaran atau pemakaian tiket langsung di produksi tanpa prosedur insiden.",
        "Sinkronkan perubahan CMS, konfigurasi tabel, tarif, dan schema ke dokumentasi.",
        "Lakukan audit berkala terhadap admin, RLS, storage policy, secret, dependency, dan log akses.",
        "Perbarui dokumen ini setelah perubahan route, status domain, integrasi, role, deployment, atau kebijakan keamanan.",
    ])
    doc.add_heading("15.1 Checklist Serah Terima", 2)
    add_table(doc, "Tabel 16. Checklist operasional", ["Item", "Status/penanggung jawab"], [
        ["Akses repository dan aturan branch", "____________________________"],
        ["Akses Vercel dan domain", "____________________________"],
        ["Akses Supabase serta backup", "____________________________"],
        ["Akses Midtrans dan webhook", "____________________________"],
        ["Akses SMTP dan alamat pengirim", "____________________________"],
        ["Akses Turnstile", "____________________________"],
        ["Daftar admin dan petugas scanner", "____________________________"],
        ["Runbook insiden dan kontak eskalasi", "____________________________"],
        ["Uji restore terakhir", "____________________________"],
        ["Uji transaksi dan scan terakhir", "____________________________"],
    ])

    # XVI
    doc.add_heading("XVI. Referensi Source Code", 1)
    add_table(doc, "Tabel 17. Lokasi implementasi penting", ["Lokasi", "Isi"], [
        ["app/components/Navbar.tsx", "Navigasi desktop/mobile, locale, akun, dan menu publik."],
        ["app/components/Dashboard/**", "Modul administrasi."],
        ["app/components/Maps/**", "Editor, preview, dan peta publik."],
        ["app/api/tickets/**", "Prepare, status, history, dan scan tiket."],
        ["app/api/midtrans/**", "Transaksi, webhook, dan reset pembayaran."],
        ["lib/tickets/**", "Keamanan status, PDF, QR, dan email tiket."],
        ["lib/security/request.ts", "Rate limiting request server."],
        ["lib/supabase/**", "Client, admin client, dan helper data."],
        ["supabase/sql/**", "Migrasi inkremental schema, constraint, RLS, dan trigger."],
        ["next.config.ts", "Redirect, image host, CSP, dan security headers."],
        ["vercel.json", "Jadwal cron cleanup draft."],
        ["PRODUCTION.md", "Runbook produksi ringkas."],
    ])
    doc.add_heading("Penutup", 1)
    doc.add_paragraph(
        "Platform menggabungkan publikasi informasi konservasi, pengelolaan data, peta interaktif, CMS bilingual, dan layanan tiket dalam satu aplikasi. Keandalan operasional bergantung pada konsistensi antara source code, schema Supabase, RLS, konfigurasi deployment, Midtrans, SMTP, serta proses backup. Gunakan dokumen ini sebagai baseline dan revisi secara terkendali ketika sistem berubah."
    )

    configure_headers(doc)
    return doc


if __name__ == "__main__":
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
